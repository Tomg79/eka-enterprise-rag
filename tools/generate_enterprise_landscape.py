"""
tools/generate_enterprise_landscape.py
======================================
Erzeugt eine realistische, ABSICHTLICH komplexe und "chaotische" Test-Firma in einen
Ziel-Ordner (das "USB-Stick"-Szenario). Deckt alle vier Quelltypen ab:

  * Dateifreigaben mit verschachtelten Ordnern + per-Datei-ACLs (inkl. KAPUTTER und
    FEHLENDER ACLs als Sicherheits-Macken),
  * mehrere SQL-Datenbanken (HR / CRM / ERP) mit teils vertraulichen Tabellen,
  * ein Object-Store-Seed (S3/MinIO) mit per-Objekt-ACLs,
  * eine SaaS-API-Quelle (Salesforce-artige JSON-Records),
  * eine Identity (policy.json/users.json) mit Rollen-Vererbung und eingebauten Macken.

Am Ende wird das Manifest `landscape.json` geschrieben -- die Discovery muss genau diese
Struktur aus dem nackten Ordner wieder rekonstruieren koennen.

Benutzung:
    python tools/generate_enterprise_landscape.py --root C:\\path\\zum\\USB\\ACME
"""

from __future__ import annotations

import argparse
import json
import os
import random
import sqlite3
import sys

# Projektwurzel importierbar machen (fuer onboarding.landscape).
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from onboarding.landscape import (
    Landscape, FileShare, Database, SqlTable, ObjectStore, ApiSource, Identity,
)

random.seed(1337)

COMPANY = "ACME Industriewerke GmbH"

# ACL-Gruppen, die in der ganzen Firma vorkommen.
GROUPS = ["PUBLIC", "FINANCE", "HR_CONFIDENTIAL", "SALES", "ENGINEERING",
          "LEGAL", "EXEC", "PROJECT_PHOENIX"]


def _w(path: str, text: str):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def _acl(path: str, groups):
    """Schreibt einen Sidecar `<datei>.acl.json` mit erlaubten Gruppen."""
    _w(path + ".acl.json", json.dumps({"groups": list(groups)}, ensure_ascii=False))


def _write_pdf(path: str, title: str, body: str, groups=None):
    """Echtes PDF (reportlab). Fehlt reportlab, Fallback als .txt."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
    except Exception:
        alt = path[:-4] + ".txt"
        _w(alt, title + "\n" + body)
        if groups:
            _acl(alt, groups)
        return
    c = canvas.Canvas(path, pagesize=A4)
    t = c.beginText(72, 780)
    for line in (title + "\n\n" + body).split("\n"):
        t.textLine(line)
    c.drawText(t)
    c.showPage()
    c.save()
    if groups:
        _acl(path, groups)


def _write_docx(path: str, title: str, body: str, groups=None):
    """Echtes DOCX (python-docx). Fehlt das Paket, Fallback als .txt."""
    os.makedirs(os.path.dirname(path), exist_ok=True)
    try:
        import docx
    except Exception:
        alt = path[:-5] + ".txt"
        _w(alt, title + "\n" + body)
        if groups:
            _acl(alt, groups)
        return
    d = docx.Document()
    d.add_heading(title, level=1)
    for para in body.split("\n"):
        d.add_paragraph(para)
    d.save(path)
    if groups:
        _acl(path, groups)


# ──────────────────────────────────────────────────────────────
# 1) Dateifreigaben (mit absichtlichen Macken)
# ──────────────────────────────────────────────────────────────
def gen_file_shares(root: str) -> list:
    shares_root = os.path.join(root, "shares")
    quirks = []

    # Finance
    _w(os.path.join(shares_root, "Finance/reports/Q1_2026_report.txt"),
       "ACME Q1 2026: Umsatz 14,2 Mio EUR, EBIT 1,8 Mio EUR.")
    _acl(os.path.join(shares_root, "Finance/reports/Q1_2026_report.txt"), ["FINANCE", "EXEC"])
    # MACKE 1: kaputter Sidecar -> Datei wird (mit sidecar-only-Reader) unsichtbar.
    _w(os.path.join(shares_root, "Finance/reports/Q2_2026_forecast.txt"),
       "VERTRAULICH: Q2-Forecast 15,0 Mio EUR.")
    _w(os.path.join(shares_root, "Finance/reports/Q2_2026_forecast.txt.acl.json"),
       "{ DAS IST KAPUTTES JSON ")
    quirks.append("Finance/reports/Q2_2026_forecast.txt: KAPUTTER ACL-Sidecar -> "
                  "mit sidecar-Reader fail-closed unsichtbar.")
    _w(os.path.join(shares_root, "Finance/budgets/budget_2026.md"),
       "# Budget 2026\nGesamtbudget 9,5 Mio EUR.")
    _acl(os.path.join(shares_root, "Finance/budgets/budget_2026.md"), ["FINANCE"])

    # HR
    _w(os.path.join(shares_root, "HR/policies/leave_policy.md"),
       "# Urlaubsrichtlinie\n30 Tage pro Jahr.")
    _acl(os.path.join(shares_root, "HR/policies/leave_policy.md"), ["PUBLIC"])  # bewusst public
    _w(os.path.join(shares_root, "HR/confidential/disciplinary_case_2026_012.txt"),
       "Abmahnung Mitarbeiter E-044 wegen wiederholter Verspaetung.")
    _acl(os.path.join(shares_root, "HR/confidential/disciplinary_case_2026_012.txt"),
         ["HR_CONFIDENTIAL"])
    # MACKE 2: gar keine ACL -> fail-closed, fuer NIEMANDEN sichtbar.
    _w(os.path.join(shares_root, "HR/confidential/salary_review_notes.txt"),
       "Gehaltsanpassungen 2026: E-001 +5%, E-002 +3%.")
    quirks.append("HR/confidential/salary_review_notes.txt: KEINE ACL -> fail-closed "
                  "fuer niemanden sichtbar (vertrauliche Gehaltsnotiz).")

    # Sales
    _w(os.path.join(shares_root, "Sales/deals/deal_CUST-097.txt"),
       "Rabatt 8% fuer CUST-097, gueltig bis Q3 2026.")
    _acl(os.path.join(shares_root, "Sales/deals/deal_CUST-097.txt"), ["SALES"])
    # MACKE 3: Sidecar schlaegt Ordner-Erwartung -> Phoenix statt SALES.
    _w(os.path.join(shares_root, "Sales/deals/project_phoenix_memo.txt"),
       "Projekt Phoenix: geheime Akquise-Strategie, nur Kernteam.")
    _acl(os.path.join(shares_root, "Sales/deals/project_phoenix_memo.txt"), ["PROJECT_PHOENIX"])
    quirks.append("Sales/deals/project_phoenix_memo.txt: per-Datei-ACL (PROJECT_PHOENIX) "
                  "schlaegt die Abteilungs-/Ordnererwartung (SALES).")

    # Engineering
    _w(os.path.join(shares_root, "Engineering/docs/system_architecture.md"),
       "# Architektur\nMikroservices, Kafka, PostgreSQL.")
    _acl(os.path.join(shares_root, "Engineering/docs/system_architecture.md"), ["ENGINEERING"])

    # Legal
    _w(os.path.join(shares_root, "Legal/contracts/contract_A_2025.txt"),
       "Rahmenvertrag mit Lieferant A, Laufzeit 3 Jahre.")
    _acl(os.path.join(shares_root, "Legal/contracts/contract_A_2025.txt"), ["LEGAL", "EXEC"])
    # MACKE 4: ACL verweist auf unbekannte Gruppe -> Discovery soll das melden.
    _w(os.path.join(shares_root, "Legal/contracts/nda_special.txt"),
       "NDA mit Partner X.")
    _acl(os.path.join(shares_root, "Legal/contracts/nda_special.txt"), ["LEGAL", "MARKETING_TYPO"])
    quirks.append("Legal/contracts/nda_special.txt: ACL referenziert unbekannte Gruppe "
                  "'MARKETING_TYPO' -> Discovery meldet verwaiste Gruppe.")

    # Verschiedene DATEIARTEN: echtes PDF + DOCX (nicht nur txt/md).
    _write_pdf(os.path.join(shares_root, "Legal/contracts/master_agreement_2026.pdf"),
               "Master Agreement 2026",
               "Rahmenvertrag mit Schluesselkunde. Volumen 4,2 Mio EUR ueber 24 Monate. "
               "Vertraulich, nur Recht und Geschaeftsfuehrung.", ["LEGAL", "EXEC"])
    _write_docx(os.path.join(shares_root, "HR/policies/employee_handbook.docx"),
                "Mitarbeiterhandbuch",
                "Arbeitszeiten, Homeoffice-Regeln, Verhaltenskodex. Gilt fuer alle Mitarbeiter.",
                ["PUBLIC"])

    return [
        FileShare(name="ACME-Fileserver", path=shares_root, acl_reader="composite",
                  note="Verschachtelte Abteilungsfreigaben mit per-Datei-ACLs (txt/md/pdf/docx)."),
    ], quirks


# ──────────────────────────────────────────────────────────────
# 2) SQL-Datenbanken (HR / CRM / ERP)
# ──────────────────────────────────────────────────────────────
def gen_databases(root: str) -> list:
    db_dir = os.path.join(root, "databases")
    os.makedirs(db_dir, exist_ok=True)
    dbs = []

    # --- HR ---
    hr = os.path.join(db_dir, "hr.db")
    if os.path.exists(hr):
        os.remove(hr)
    con = sqlite3.connect(hr); c = con.cursor()
    c.execute("CREATE TABLE T_EMPLOYEES (EMP_ID TEXT PRIMARY KEY, EMP_NAME TEXT, DEPT TEXT, EMAIL TEXT)")
    c.execute("CREATE TABLE T_SALARIES (EMP_ID TEXT PRIMARY KEY, SALARY_EUR REAL, BONUS_EUR REAL)")
    c.execute("CREATE TABLE T_DEPARTMENTS (DEPT TEXT PRIMARY KEY, HEAD_EMP_ID TEXT, COST_CENTER TEXT)")
    depts = ["Finance", "HR", "Sales", "Engineering", "Legal", "IT"]
    emps = []
    for i in range(1, 61):
        eid = f"E-{i:03d}"
        dept = random.choice(depts)
        emps.append((eid, f"Mitarbeiter_{i}", dept, f"user{i}@acme.example"))
    c.executemany("INSERT INTO T_EMPLOYEES VALUES (?,?,?,?)", emps)
    c.executemany("INSERT INTO T_SALARIES VALUES (?,?,?)",
                  [(e[0], round(random.uniform(45000, 160000), 2), round(random.uniform(0, 20000), 2)) for e in emps])
    c.executemany("INSERT INTO T_DEPARTMENTS VALUES (?,?,?)",
                  [(d, random.choice(emps)[0], f"CC-{1000+i}") for i, d in enumerate(depts)])
    con.commit(); con.close()
    dbs.append(Database(
        name="hr", connection_string=f"sqlite:///{hr}",
        tables=[
            SqlTable("T_EMPLOYEES", "EMP_ID", ["EMP_ID", "EMP_NAME", "DEPT", "EMAIL"]),
            SqlTable("T_SALARIES", "EMP_ID", ["EMP_ID", "SALARY_EUR", "BONUS_EUR"], sensitive=True),
            SqlTable("T_DEPARTMENTS", "DEPT", ["DEPT", "HEAD_EMP_ID", "COST_CENTER"]),
        ],
        note="Personaldatenbank. T_SALARIES ist vertraulich (nur HR/EXEC)."))

    # --- CRM ---
    crm = os.path.join(db_dir, "crm.db")
    if os.path.exists(crm):
        os.remove(crm)
    con = sqlite3.connect(crm); c = con.cursor()
    c.execute("CREATE TABLE T_CUSTOMERS (CUST_ID TEXT PRIMARY KEY, NAME TEXT, TIER TEXT, ACCOUNT_MGR TEXT)")
    c.execute("CREATE TABLE T_CONTACTS (CONTACT_ID TEXT PRIMARY KEY, CUST_ID TEXT, NAME TEXT, EMAIL TEXT)")
    custs = [(f"CUST-{i:03d}", f"Kunde {i}", random.choice(["A", "B", "C"]), f"E-{random.randint(1,60):03d}") for i in range(1, 121)]
    c.executemany("INSERT INTO T_CUSTOMERS VALUES (?,?,?,?)", custs)
    c.executemany("INSERT INTO T_CONTACTS VALUES (?,?,?,?)",
                  [(f"CT-{i:04d}", f"CUST-{random.randint(1,120):03d}", f"Kontakt {i}", f"kontakt{i}@kunde.example") for i in range(1, 200)])
    con.commit(); con.close()
    dbs.append(Database(
        name="crm", connection_string=f"sqlite:///{crm}",
        tables=[
            SqlTable("T_CUSTOMERS", "CUST_ID", ["CUST_ID", "NAME", "TIER", "ACCOUNT_MGR"]),
            SqlTable("T_CONTACTS", "CONTACT_ID", ["CONTACT_ID", "CUST_ID", "NAME", "EMAIL"]),
        ],
        note="CRM. Sichtbar fuer Sales/EXEC."))

    # --- ERP ---
    erp = os.path.join(db_dir, "erp.db")
    if os.path.exists(erp):
        os.remove(erp)
    con = sqlite3.connect(erp); c = con.cursor()
    c.execute("CREATE TABLE T_SALES_ORDERS (ORDER_ID TEXT PRIMARY KEY, CUST_ID TEXT, MAT_NR TEXT, QTY INTEGER, NET_EUR REAL, STATUS TEXT)")
    c.execute("CREATE TABLE T_MATERIALS (MAT_NR TEXT PRIMARY KEY, MAT_DESC TEXT, UNIT TEXT, PRICE_EUR REAL)")
    c.execute("CREATE TABLE T_INVENTORY (MAT_NR TEXT PRIMARY KEY, STOCK_QTY INTEGER, WAREHOUSE TEXT)")
    mats = [(f"MAT-{i:03d}", f"Material {i}", random.choice(["Stk", "kg", "m"]), round(random.uniform(5, 900), 2)) for i in range(1, 101)]
    c.executemany("INSERT INTO T_MATERIALS VALUES (?,?,?,?)", mats)
    c.executemany("INSERT INTO T_INVENTORY VALUES (?,?,?)",
                  [(m[0], random.randint(0, 5000), random.choice(["WH-1", "WH-2", "WH-3"])) for m in mats])
    c.executemany("INSERT INTO T_SALES_ORDERS VALUES (?,?,?,?,?,?)",
                  [(f"SO-{1000+i}", f"CUST-{random.randint(1,120):03d}", f"MAT-{random.randint(1,100):03d}",
                    random.randint(1, 300), round(random.uniform(10, 9000), 2),
                    random.choice(["OPEN", "SHIPPED", "INVOICED", "CANCELLED"])) for i in range(1, 161)])
    con.commit(); con.close()
    dbs.append(Database(
        name="erp", connection_string=f"sqlite:///{erp}",
        tables=[
            SqlTable("T_SALES_ORDERS", "ORDER_ID", ["ORDER_ID", "CUST_ID", "MAT_NR", "QTY", "NET_EUR", "STATUS"]),
            SqlTable("T_MATERIALS", "MAT_NR", ["MAT_NR", "MAT_DESC", "UNIT", "PRICE_EUR"]),
            SqlTable("T_INVENTORY", "MAT_NR", ["MAT_NR", "STOCK_QTY", "WAREHOUSE"]),
        ],
        note="ERP. Auftraege/Material/Bestand."))

    # Zweiter DB-TYP: Postgres-Data-Warehouse (nur Connection-String; der Server laeuft
    # separat, ist NICHT Teil der Demo-Dateien). Zeigt, dass der SQL-Connector beliebige
    # DBMS kann, nicht nur SQLite. Credentials gehoeren in eine Env-Variable, nicht hierher.
    dbs.append(Database(
        name="warehouse",
        connection_string="postgresql://eka_ro:CHANGE_ME@db.acme.local:5432/warehouse",
        tables=[
            SqlTable("FACT_SALES", "SALE_ID",
                     ["SALE_ID", "CUST_ID", "AMOUNT_EUR", "REGION", "SALE_DATE"]),
            SqlTable("DIM_CUSTOMER", "CUST_ID", ["CUST_ID", "NAME", "SEGMENT"]),
        ],
        note="Postgres-DWH (anderer DB-Typ als die SQLite-Quellen)."))

    return dbs


# ──────────────────────────────────────────────────────────────
# 3) Object-Store-Seed (S3/MinIO)
# ──────────────────────────────────────────────────────────────
def gen_object_store(root: str) -> list:
    seed = os.path.join(root, "object_store", "corp-archive")
    _w(os.path.join(seed, "legal/old_contract_2019.txt"), "Altvertrag 2019, archiviert.")
    _acl(os.path.join(seed, "legal/old_contract_2019.txt"), ["LEGAL"])
    _w(os.path.join(seed, "sales/archived_deal_2024.txt"), "Abgeschlossener Deal 2024.")
    _acl(os.path.join(seed, "sales/archived_deal_2024.txt"), ["SALES"])
    return [ObjectStore(
        name="corp-archive", endpoint_url="", bucket="corp-archive", prefix="",
        local_seed_path=seed, acl_reader="composite",
        credentials_env="MINIO",
        note="Lokaler Seed-Ordner. Fuer echtes S3/MinIO: endpoint_url+Credentials setzen, "
             "Seed hochladen (deploy/minio_seed.sh)."),
    ]


# ──────────────────────────────────────────────────────────────
# 4) SaaS-API (Salesforce-artig)
# ──────────────────────────────────────────────────────────────
def gen_apis(root: str) -> list:
    api_dir = os.path.join(root, "apis", "salesforce")
    os.makedirs(api_dir, exist_ok=True)
    for i in range(1, 31):
        rec = {"customer_id": f"CUST-{i:03d}", "account_manager": f"E-{random.randint(1,60):03d}",
               "status": random.choice(["active", "churned", "prospect"]),
               "arr_eur": random.randint(1000, 500000)}
        _w(os.path.join(api_dir, f"CUST-{i:03d}.json"), json.dumps(rec, ensure_ascii=False))
    return [ApiSource(name="salesforce", kind="salesforce_json", path=api_dir,
                      note="JSON-Records; Scope ueber PolicyEngine (none/all/list).")]


# ──────────────────────────────────────────────────────────────
# 5) Identity (Rollen/User mit Macken)
# ──────────────────────────────────────────────────────────────
def gen_identity(root: str) -> Identity:
    idir = os.path.join(root, "identity")
    os.makedirs(idir, exist_ok=True)
    policy = {
        "_comment": "ACME RBAC: Abteilungsrollen + EXEC erbt alle. Enthaelt absichtliche Macken.",
        "roles": {
            "base": {"inherits": [], "groups": ["PUBLIC"], "dms_tags": ["PUBLIC"],
                     "sap_tables": [], "sql_tables": [], "salesforce": {"mode": "none"}},
            "finance": {"inherits": ["base"], "groups": ["FINANCE"], "dms_tags": ["FINANCE"],
                        "sql_tables": ["hr.T_DEPARTMENTS"], "salesforce": {"mode": "none"}},
            "hr": {"inherits": ["base"], "groups": ["HR_CONFIDENTIAL"], "dms_tags": ["HR_CONFIDENTIAL"],
                   "sql_tables": ["hr.T_EMPLOYEES", "hr.T_SALARIES", "hr.T_DEPARTMENTS"],
                   "salesforce": {"mode": "none"}},
            "sales": {"inherits": ["base"], "groups": ["SALES"], "dms_tags": ["SALES"],
                      "sql_tables": ["crm.T_CUSTOMERS", "crm.T_CONTACTS", "erp.T_SALES_ORDERS", "erp.T_MATERIALS"],
                      "salesforce": {"mode": "all"}},
            "engineering": {"inherits": ["base"], "groups": ["ENGINEERING"], "dms_tags": ["ENGINEERING"],
                            "sql_tables": ["erp.T_MATERIALS", "erp.T_INVENTORY"], "salesforce": {"mode": "none"}},
            "legal": {"inherits": ["base"], "groups": ["LEGAL"], "dms_tags": ["LEGAL"],
                      "sql_tables": [], "salesforce": {"mode": "none"}},
            "exec": {"inherits": ["finance", "hr", "sales", "engineering", "legal"],
                     "groups": ["EXEC"], "dms_tags": [], "sql_tables": [], "salesforce": {"mode": "all"}},
            "project_phoenix": {"inherits": [], "groups": ["PROJECT_PHOENIX"], "dms_tags": [],
                                "sql_tables": [], "salesforce": {"mode": "none"}},
            # MACKE 5: Rolle referenziert eine nicht existierende Gruppe.
            "marketing": {"inherits": ["base"], "groups": ["MARKETING"], "dms_tags": [],
                          "sql_tables": [], "salesforce": {"mode": "none"}},
        },
    }
    with open(os.path.join(idir, "policy.json"), "w", encoding="utf-8") as f:
        json.dump(policy, f, indent=2, ensure_ascii=False)

    users = {
        "_comment": "ACME-Mitarbeiter. Enthaelt Multi-Rollen, Konflikte und einen Geist-User.",
        "users": {
            "ceo":        {"display": "Clara Berg (CEO)",            "roles": ["exec"]},
            "cfo":        {"display": "Frank Mertens (CFO)",         "roles": ["finance", "exec"]},
            "hr_head":    {"display": "Petra Lang (HR-Leitung)",     "roles": ["hr"]},
            "sales_alice":{"display": "Alice Sommer (Sales+Phoenix)","roles": ["sales", "project_phoenix"]},
            "sales_bob":  {"display": "Bob Klein (Sales)",           "roles": ["sales"]},
            "eng_lead":   {"display": "Dr. Eva Roth (Eng-Lead)",     "roles": ["engineering"]},
            "legal_max":  {"display": "Max Vogt (Justiziar)",        "roles": ["legal"]},
            "intern_ida": {"display": "Ida Neu (Praktikantin)",      "roles": ["base"]},
            "erin":       {"display": "Erin Falk (Sales+HR)",        "roles": ["sales", "hr"]},
            "conflict_carl": {"display": "Carl Fuchs (Finance+Legal)", "roles": ["finance", "legal"]},
            # MACKE 6: User mit unbekannter Rolle -> faellt auf keine Rechte (fail-closed).
            "ghost_gwen": {"display": "Gwen Geist (?)", "roles": ["nonexistent_role"]},
        },
    }
    with open(os.path.join(idir, "users.json"), "w", encoding="utf-8") as f:
        json.dump(users, f, indent=2, ensure_ascii=False)

    return Identity(policy_path=os.path.join(idir, "policy.json"),
                    users_path=os.path.join(idir, "users.json"))


def main():
    ap = argparse.ArgumentParser(description="Erzeugt eine komplexe Test-Unternehmens-Landscape.")
    ap.add_argument("--root", default="enterprise_demo", help="Zielordner (USB-Stick-Verzeichnis).")
    args = ap.parse_args()
    root = os.path.abspath(args.root)
    os.makedirs(root, exist_ok=True)

    print(f"Erzeuge Test-Firma '{COMPANY}' in: {root}")
    file_shares, quirks = gen_file_shares(root)
    databases = gen_databases(root)
    object_stores = gen_object_store(root)
    apis = gen_apis(root)
    identity = gen_identity(root)

    notes = ["Synthetische Demo-Firma mit absichtlichen Sicherheits-Macken:"] + \
            [f"  - {q}" for q in quirks] + \
            ["  - identity: Rolle 'marketing' referenziert unbekannte Gruppe 'MARKETING_TYPO'.",
             "  - identity: User 'ghost_gwen' hat eine unbekannte Rolle -> keine Rechte."]

    ls = Landscape(company=COMPANY, file_shares=file_shares, databases=databases,
                   object_stores=object_stores, apis=apis, identity=identity, notes=notes)
    ls.save(os.path.join(root, "landscape.json"))

    readme = (f"# {COMPANY} -- synthetische Test-Landscape\n\n"
              "Erzeugt von tools/generate_enterprise_landscape.py.\n\n"
              "Inhalt: shares/ (txt/md/pdf/docx + ACLs), databases/ (sqlite hr/crm/erp + "
              "Postgres-DWH-Eintrag), object_store/ (S3-Seed), apis/ (JSON), identity/ "
              "(policy+users), landscape.json (Manifest).\n")
    _w(os.path.join(root, "README_LANDSCAPE.md"), readme)

    print(f"Fertig. {len(file_shares)} Freigabe(n), {len(databases)} DB(s), "
          f"{len(object_stores)} Object-Store(s), {len(apis)} API(s).")
    print(f"Manifest: {os.path.join(root, 'landscape.json')}")


if __name__ == "__main__":
    main()
