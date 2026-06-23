"""
connectors/filesystem.py -- Filesystem-Connector (Phase 2, erweitert).

Liest Dokumente + Rechte-/Config-Dateien unter einem Wurzelverzeichnis und ermittelt pro
Datei die erlaubten Gruppen via AclReader. Sidecars (*.acl.json) werden uebersprungen.

Formate (Standard-Libs/lazy): txt, md, csv, sql, yaml/yml, graphql, http, json (Text);
pdf (pypdf/PyPDF2); docx (python-docx); xlsx (openpyxl). So werden auch Rechte-"Sprachen"
wie SQL-GRANT, OpenAPI x-permissions, Kubernetes-RBAC und GraphQL-@auth durchsuchbar +
RBAC-gefiltert. Fehlt eine Lib, wird die Datei uebersprungen (fail-safe).

Hinweis: Veraltete/widerrufene Dateien werden NICHT mehr stillschweigend ausgeschlossen --
sie sind Teil der Rechte-Historie (z.B. Widerrufs-Beschluesse) und werden RBAC-gefiltert
indexiert, damit der Assistent den aktuellen Stand herleiten kann.
"""

from __future__ import annotations

import logging
import os
from typing import Iterator

from .base import Connector, ConnectorDocument
from .acl_readers import AclReader

logger = logging.getLogger(__name__)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _read_pdf(path: str) -> str:
    try:
        try:
            from pypdf import PdfReader
        except Exception:
            from PyPDF2 import PdfReader
    except Exception as e:
        logger.error("FS: PDF-Lib fehlt (%s) -> '%s' uebersprungen.", e, path)
        return ""
    reader = PdfReader(path)
    out = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(out)


def _read_docx(path: str) -> str:
    try:
        import docx
    except Exception as e:
        logger.error("FS: python-docx fehlt (%s) -> '%s' uebersprungen.", e, path)
        return ""
    d = docx.Document(path)
    return "\n".join(p.text for p in d.paragraphs)


def _read_xlsx(path: str) -> str:
    try:
        import openpyxl
    except Exception as e:
        logger.error("FS: openpyxl fehlt (%s) -> '%s' uebersprungen.", e, path)
        return ""
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# Blatt: {ws.title}")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            out.append("; ".join("" if c is None else str(c) for c in row))
            if i > 500:
                break
    return "\n".join(out)


# Endung -> Reader. Text-/Config-Formate teilen sich _read_txt.
_READERS = {
    ".txt": _read_txt, ".md": _read_txt, ".csv": _read_txt, ".sql": _read_txt,
    ".yaml": _read_txt, ".yml": _read_txt, ".graphql": _read_txt, ".http": _read_txt,
    ".json": _read_txt,
    ".pdf": _read_pdf, ".docx": _read_docx, ".xlsx": _read_xlsx,
}

# Endungen, die NIE indexiert werden (Secrets/Metadaten/Muell).
_SKIP_SUFFIX = (".acl.json", ".bak", ".lock", ".tmp", ".bin", ".ds_store",
                ".ini", ".env")  # .ini/.env = Zugangsdaten -> NIE indexieren


class FilesystemConnector(Connector):
    name = "filesystem"
    TEXT_EXTENSIONS = set(_READERS.keys())

    def __init__(self, root_dir: str, acl_reader: AclReader, recursive: bool = True):
        self.root_dir = root_dir
        self.acl_reader = acl_reader
        self.recursive = recursive

    def iter_documents(self) -> Iterator[ConnectorDocument]:
        for dirpath, _dirs, files in os.walk(self.root_dir):
            for fn in sorted(files):
                low = fn.lower()
                if low.endswith(_SKIP_SUFFIX):
                    continue
                ext = os.path.splitext(fn)[1].lower()
                reader = _READERS.get(ext)
                if reader is None:
                    continue
                path = os.path.join(dirpath, fn)
                try:
                    text = reader(path)
                except Exception as e:
                    logger.error("FS: '%s' nicht lesbar (%s) -> uebersprungen.", path, e)
                    continue
                if not (text or "").strip():
                    continue
                groups = self.acl_reader.groups_for(path)
                if not groups:
                    logger.warning("FS: '%s' ohne erlaubte Gruppen (ACL leer) -> "
                                   "indexiert, fuer niemanden sichtbar.", path)
                yield ConnectorDocument(
                    doc_id=fn, text=text, acl_groups=groups,
                    metadata={"file_name": fn, "source_type": "document", "path": path},
                )
            if not self.recursive:
                break
